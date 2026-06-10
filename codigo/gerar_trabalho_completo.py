import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ════════════════════════════════════════════════════════════════════
# Inicializa o documento e copia as configurações de página
# ════════════════════════════════════════════════════════════════════
template = Document(r'c:\andd\ap2_analise_de_dados\4_Referencias_Bibliograficas.docx')
doc = Document()

# Configura as margens e dimensões a partir do template
for sec_src, sec_dst in zip(template.sections, doc.sections):
    sec_dst.page_width = sec_src.page_width
    sec_dst.page_height = sec_src.page_height
    sec_dst.left_margin = sec_src.left_margin
    sec_dst.right_margin = sec_src.right_margin
    sec_dst.top_margin = sec_src.top_margin
    sec_dst.bottom_margin = sec_src.bottom_margin

# ════════════════════════════════════════════════════════════════════
# Paleta de Cores e Estilos Corporativos
# ════════════════════════════════════════════════════════════════════
FONT_NAME = "Arial"
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)   # #1F3864
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_BODY = RGBColor(0x26, 0x26, 0x26)   # #262626
GRAY_TEXT = RGBColor(0x59, 0x59, 0x59)   # #595959

def add_heading1(doc, text):
    """Título de Primeiro Nível (Seção Principal) com preenchimento sutil."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:color="1F3864"/>')
    pPr.append(shd)
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="360"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200" w:right="200"/>')
    pPr.append(ind)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = WHITE
    return p

def add_heading2(doc, text):
    """Subtítulo de Segundo Nível em Azul Escuro sem preenchimento."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:before="240"/>')
    pPr.append(spacing)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    return p

def add_body_paragraph(doc, text, bold_prefix=None, align_justify=True):
    """Parágrafo regular justificado com fonte Arial 11."""
    p = doc.add_paragraph()
    if align_justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="160" w:before="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
        r_prefix.font.name = FONT_NAME
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = TEXT_BODY
        
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_BODY
    return p

def add_quote_box(doc, text):
    """Bloco de citação ou destaque com recuo e borda esquerda cinza."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="100"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:right="360"/>')
    pPr.append(ind)
    pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="12" w:color="1F3864"/></w:pBdr>')
    pPr.append(pbdr)
    
    run = p.add_run(text)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_TEXT
    return p

def create_table(doc, rows, cols, data, col_widths=None):
    """Gera tabelas elegantes e limpas."""
    table_shape = doc.add_table(rows, cols)
    table = table_shape
    
    # Define largura das colunas
    if col_widths:
        for r in table.rows:
            for idx, w in enumerate(col_widths):
                r.cells[idx].width = w
                
    for r_idx in range(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            
            # Limpa o texto padrão e adiciona formatado
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._element.get_or_add_pPr()
            spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="60" w:before="60"/>')
            pPr.append(spacing)
            
            run = p.add_run(str(data[r_idx][c_idx]))
            run.font.name = FONT_NAME
            run.font.size = Pt(9.5)
            run.font.color.rgb = WHITE if r_idx == 0 else TEXT_BODY
            run.bold = True if (r_idx == 0 or c_idx == 0) else False
            
            tcPr = cell._tc.get_or_add_tcPr()
            # Estiliza o preenchimento de fundo da célula via XML shading
            if r_idx == 0:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:color="1F3864"/>')
            else:
                if r_idx % 2 == 1:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF" w:color="FFFFFF"/>')
                else:
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9" w:color="F1F5F9"/>')
            tcPr.append(shd)
            
            # Adiciona bordas finas horizontais às células
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
                f'  <w:left w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)
            
    # Linha vazia após a tabela
    p_after = doc.add_paragraph()
    pPr_after = p_after._element.get_or_add_pPr()
    spacing_after = parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:before="0"/>')
    pPr_after.append(spacing_after)

def add_reference_abnt(doc, author, title, rest, url=None):
    """Adiciona entrada de referência bibliográfica no padrão ABNT."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="160" w:before="0"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="720"/>')
    pPr.append(ind)
    
    r_author = p.add_run(author)
    r_author.bold = True
    r_author.font.name = FONT_NAME
    r_author.font.size = Pt(11)
    r_author.font.color.rgb = TEXT_BODY
    
    r_space = p.add_run(" ")
    r_space.font.name = FONT_NAME
    r_space.font.size = Pt(11)
    
    r_title = p.add_run(title)
    r_title.bold = True
    r_title.font.name = FONT_NAME
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = TEXT_BODY
    
    r_rest = p.add_run(rest)
    r_rest.font.name = FONT_NAME
    r_rest.font.size = Pt(11)
    r_rest.font.color.rgb = TEXT_BODY
    
    if url:
        r_url_label = p.add_run(" Disponível em: ")
        r_url_label.font.name = FONT_NAME
        r_url_label.font.size = Pt(11)
        r_url_label.font.color.rgb = TEXT_BODY
        
        part = p.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
        new_run = parse_xml(
            f'<w:r {nsdecls("w")}>'
            f'  <w:rPr>'
            f'    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial" w:eastAsia="Arial"/>'
            f'    <w:color w:val="1155CC"/>'
            f'    <w:u w:val="single"/>'
            f'    <w:sz w:val="22"/>'
            f'  </w:rPr>'
            f'  <w:t>{url}</w:t>'
            f'</w:r>'
        )
        hyperlink.append(new_run)
        p._element.append(hyperlink)
        
        r_after = p.add_run(". Acesso em: 10 jun. 2026.")
        r_after.font.name = FONT_NAME
        r_after.font.size = Pt(11)
        r_after.font.color.rgb = TEXT_BODY

def add_empty_line(doc):
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)

# Remove parágrafo inicial vazio
if doc.paragraphs:
    doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

# ════════════════════════════════════════════════════════════════════
# 📄 ESTRUTURAÇÃO DO TRABALHO
# ════════════════════════════════════════════════════════════════════

# ── PÁGINA 1: CAPA ──
p_institucional = doc.add_paragraph()
p_institucional.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_institucional._element.get_or_add_pPr()
pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="1200" w:before="0"/>'))
r_inst = p_institucional.add_run("CENTRO DE ESTUDOS AVANÇADOS EM FINANÇAS CORPORATIVAS\nCURSO DE ANÁLISE QUANTITATIVA E VALUATION")
r_inst.bold = True
r_inst.font.name = FONT_NAME
r_inst.font.size = Pt(12)
r_inst.font.color.rgb = DARK_BLUE

p_titulo = doc.add_paragraph()
p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_t = p_titulo._element.get_or_add_pPr()
pPr_t.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="400" w:before="800"/>'))
r_title = p_titulo.add_run("ESTUDO COMPARATIVO MULTICRITÉRIO DE DESEMPENHO FINANCEIRO E RISCO DE MERCADO")
r_title.bold = True
r_title.font.name = FONT_NAME
r_title.font.size = Pt(18)
r_title.font.color.rgb = DARK_BLUE

p_subtitulo = doc.add_paragraph()
p_subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_sub = p_subtitulo._element.get_or_add_pPr()
pPr_sub.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="2400" w:before="0"/>'))
r_sub = p_subtitulo.add_run("Análise Integrada de Demonstrações Contábeis, Economic Value Added (EVA), Performance de Ações e Simulação de Monte Carlo de ABEV3, ASAI3 e AMER3 (2021-2025)")
r_sub.italic = True
r_sub.font.name = FONT_NAME
r_sub.font.size = Pt(12)
r_sub.font.color.rgb = GRAY_TEXT

p_autoria = doc.add_paragraph()
p_autoria.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_aut = p_autoria._element.get_or_add_pPr()
pPr_aut.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="1200" w:before="0"/>'))
r_aut = p_autoria.add_run("Elaborado por:\nGRUPO DE ESTUDOS E PESQUISAS EM MERCADO DE CAPITAIS\n\nOrientador: Prof. Dr. em Finanças Quantitativas")
r_aut.bold = True
r_aut.font.name = FONT_NAME
r_aut.font.size = Pt(11)
r_aut.font.color.rgb = TEXT_BODY

p_data = doc.add_paragraph()
p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr_data = p_data._element.get_or_add_pPr()
pPr_data.append(parse_xml(f'<w:spacing {nsdecls("w")} w:after="0" w:before="800"/>'))
r_data = p_data.add_run("SÃO PAULO\nJUNHO / 2026")
r_data.bold = True
r_data.font.name = FONT_NAME
r_data.font.size = Pt(10)
r_data.font.color.rgb = GRAY_TEXT

doc.add_page_break()

# ── PÁGINA 2: INTRODUÇÃO E CENÁRIO MACRO ──
add_heading1(doc, "1. Introdução e Dinâmica Macroeconômica (2021-2025)")

add_body_paragraph(doc,
    "Este trabalho apresenta uma análise financeira multicritério abrangendo o período de 2021 a 2025, integrando a "
    "avaliação contábil clássica, a mensuração de criação de valor residual por meio do Economic Value Added (EVA), "
    "a estatística de retornos e volatilidade das ações no mercado secundário e a modelagem estocástica de risco futuro "
    "via Simulação de Monte Carlo. O escopo do estudo compreende três grandes empresas listadas na B3 que operam com "
    "diferentes modelos de negócios, custos de capital e estratégias de financiamento: Ambev S.A. (ABEV3), Assaí Atacadista "
    "(ASAI3) e Lojas Americanas S.A. (AMER3)."
)

add_body_paragraph(doc,
    "Ao longo dos cinco anos em análise, o cenário macroeconômico brasileiro impôs severos desafios às corporações de varejo "
    "e consumo. O período pós-pandemia caracterizou-se pela reaceleração inflacionária, o que obrigou o Banco Central do "
    "Brasil a elevar a taxa de juros básica (Selic) de patamares mínimos históricos de 2,0% para a casa de dois dígitos, "
    "fechando o ano de 2025 com uma Selic média anual de 10,75%. Esse movimento encareceu de forma generalizada o custo de "
    "captação de recursos (taxa livre de risco), encarecendo o carrego de capital de giro e as dívidas bancárias, além de "
    "contrair a renda real das famílias brasileiras e elevar o nível de endividamento dos consumidores."
)

add_body_paragraph(doc,
    "Nesse contexto de juros elevados e aversão ao risco no mercado de capitais, a estrutura de capital e a alavancagem financeira "
    "tornaram-se os principais divisores de águas entre a solidez corporativa e a ruína. Empresas com balanços blindados e "
    "geração de caixa operacional previsível foram altamente valorizadas, enquanto corporações excessivamente alavancadas "
    "ou envolvidas em escândalos corporativos foram duramente corrigidas pelas forças de mercado."
)

# ── PÁGINA 3: PERFIS CORPORATIVOS ──
add_heading1(doc, "2. Perfis Corporativos das Empresas Analisadas")

add_heading2(doc, "2.1. Ambev S.A. (ABEV3) — O Gigante de Consumo Defensivo")
add_body_paragraph(doc,
    "A Ambev S.A. é a maior fabricante de cervejas e bebidas da América Latina, com uma participação de mercado dominante no "
    "território brasileiro. Seu modelo de negócios apoia-se em uma fantástica escala de fabricação e em uma rede de distribuição "
    "própria com capilaridade sem paralelos no país, alcançando praticamente 100% do varejo físico de bebidas. A empresa opera com "
    "marcas consolidadas e tem direcionado esforços para a digitalização de seus canais B2B (plataforma Bees) e B2C (aplicativo "
    "Zé Delivery). Do ponto de vista de finanças corporativas, a Ambev caracteriza-se por um modelo de consumo não cíclico defensivo, "
    "geração robusta de fluxo de caixa livre e uma estrutura patrimonial extremamente conservadora, operando historicamente livre "
    "de endividamento líquido (com caixa líquido excedente). Os principais riscos estratégicos no horizonte envolvem a Reforma "
    "Tributária nacional (incidência de imposto seletivo) e potenciais mudanças regulatórias na dedutibilidade fiscal do Juros "
    "sobre Capital Próprio (JCP)."
)

add_heading2(doc, "2.2. Assaí Atacadista (ASAI3) — Expansão sob Alavancagem no Atacarejo")
add_body_paragraph(doc,
    "O Assaí Atacadista opera no modelo Cash & Carry (atacarejo), focado em vendas em grande escala com margens de comercialização "
    "muito estreitas e altíssimo giro de estoques. O atacarejo consolidou-se como o principal canal de abastecimento das famílias "
    "brasileiras em momentos de inflação de alimentos. Em 2021, a empresa realizou sua cisão societária (spin-off) do Grupo Pão "
    "de Açúcar (GPA), tornando-se uma corporation pulverizada na B3. Visando acelerar seu crescimento e dominar pontos de venda premium "
    "nos grandes centros urbanos, o Assaí realizou a aquisição de 70 lojas do Extra Hipermercado. No entanto, o capex volumoso para a "
    "conversão e compra desses ativos foi financiado por captação massiva de dívidas no mercado bancário. A permanência da taxa Selic "
    "elevada encareceu severamente as despesas financeiras líquidas da empresa, consumindo a rentabilidade líquida do acionista."
)

add_heading2(doc, "2.3. Lojas Americanas S.A. (AMER3) — O Colapso Administrativo-Contábil")
add_body_paragraph(doc,
    "Fundada em 1929, as Lojas Americanas constituem uma das marcas de varejo físico e eletrônico mais tradicionais e conhecidas "
    "do país. Historicamente focada em utilidades, conveniências e eletrônicos, a empresa integrou sua rede de lojas físicas com o "
    "e-commerce e marketplace digital de larga escala. No entanto, a trajetória da companhia sofreu um ponto de inflexão catastrófico "
    "em janeiro de 2023, com a descoberta de fraude contábil bilionária de aproximadamente R$ 20 bilhões omitidos em passivos de risco "
    "sacado (forfaiting). A revelação do rombo contábil congelou as linhas de crédito bancário da empresa, forçando o pedido imediato "
    "de Recuperação Judicial. O plano de sobrevivência homologado envolveu um aumento de capital de R$ 12 bilhões liderado pelos acionistas "
    "de referência (grupamento 3G Capital) e conversão em larga escala de passivos de credores em novas ações. Em agosto de 2024, a B3 "
    "executou o grupamento de ações de 100 para 1 para evitar a listagem continuada do ativo como penny stock."
)

# ── PÁGINA 4: ANÁLISE CONTÁBIL CONSOLIDADA ──
add_heading1(doc, "3. Análise Contábil e Estrutura de Balanços (Exercício de 2025)")

add_body_paragraph(doc,
    "As Demonstrações Financeiras Padronizadas (DFPs) auditadas do exercício encerrado em 31 de dezembro de 2025 revelam a real "
    "condição de eficiência e liquidez de cada empresa. O quadro contábil a seguir resume as principais linhas de resultado, as margens "
    "operacionais e a estrutura de endividamento líquido."
)

# Dados da Tabela Contábil
data_contabil = [
    ["Grupo / Indicador Financeiro (2025)", "ABEV3 (Ambev)", "AMER3 (Americanas)", "ASAI3 (Assaí)"],
    ["Receita Líquida", "R$ 82.593 mi", "R$ 12.328 mi", "R$ 71.519 mi"],
    ["Lucro Bruto", "R$ 42.474 mi", "R$ 3.124 mi", "R$ 12.066 mi"],
    ["Margem Bruta (%)", "51,42%", "25,34%", "16,87%"],
    ["EBITDA", "R$ 28.324 mi", "R$ 684 mi", "R$ 3.623 mi"],
    ["Margem EBITDA (%)", "34,29%", "5,55%", "5,07%"],
    ["EBIT (Lucro Operacional)", "R$ 21.917 mi", "R$ 331 mi", "R$ 3.623 mi"],
    ["Margem EBIT (%)", "26,54%", "2,69%", "5,07%"],
    ["Lucro Líquido", "R$ 14.968 mi", "-R$ 271 mi", "R$ 1.023 mi"],
    ["Margem Líquida (%)", "18,12%", "-2,20%", "1,43%"],
    ["Patrimônio Líquido (PL)", "R$ 82.599 mi", "R$ 5.372 mi", "R$ 4.705 mi"],
    ["Dívida Líquida", "-R$ 8.440 mi (Caixa)", "R$ 5.772 mi", "R$ 21.319 mi"]
]

col_widths = [Inches(3.4), Inches(2.2), Inches(2.2), Inches(2.2)]
create_table(doc, 12, 4, data_contabil, col_widths)

add_body_paragraph(doc,
    "A análise da eficiência em nível bruto revela que a Ambev opera com alta rentabilidade (margem bruta de 51,42%), impulsionada "
    "pelo pricing power de suas marcas premium e baixíssimos custos diretos de fabricação decorrentes de ganhos de escala. O Assaí, "
    "como é típico do modelo de varejo alimentar de alto volume, apresenta uma margem bruta reduzida de 16,87%, exigindo altíssimo giro "
    "para cobrir suas despesas. A Americanas apresenta margem bruta de 25,34%, mas sua receita encolhida aponta para perda de escala comercial "
    "e dificuldades na negociação com fornecedores pós-fraude."
)

add_body_paragraph(doc,
    "A disparidade operacional é ainda mais marcante no EBIT. A Ambev retém 26,54% de sua receita sob forma de lucro operacional "
    "puro (R$ 21.917 milhões), demonstrando alto retorno operacional. O Assaí atinge R$ 3.623 milhões de EBIT (margem de 5,07%), mas "
    "seu faturamento é corroído pelas despesas com juros, reduzindo seu lucro líquido final a R$ 1.023 milhões (margem líquida estreita "
    "de 1,43%). Por fim, a Americanas reporta EBIT de apenas R$ 331 milhões (margem operacional de 2,69%), insuficiente para arcar com os "
    "encargos financeiros da recuperação judicial, culminando em prejuízo líquido de R$ 271 milhões."
)

add_quote_box(doc,
    "Destaque Patrimonial: A Ambev detém R$ 8.440 milhões de caixa líquido de obrigações, blindando o negócio contra ciclos de crédito. "
    "Em oposição, a Dívida Líquida do Assaí de R$ 21.319 milhões representa 4,53 vezes o seu Patrimônio Líquido contábil, evidenciando "
    "extremo risco patrimonial sob juros em dois dígitos."
)

# ── PÁGINA 5: CRIAÇÃO DE VALOR ECONÔMICO (EVA) ──
add_heading1(doc, "4. Criação de Valor Econômico (Economic Value Added — EVA)")

add_body_paragraph(doc,
    "O lucro contábil reportado na DRE é insuficiente para atestar a eficiência alocativa da gestão. Para avaliar a real criação de "
    "riqueza residual, aplicou-se a modelagem de Stern Stewart baseada no Economic Value Added (EVA). O cálculo deduz o custo de "
    "oportunidade dos sócios e credores (WACC) do lucro operacional pós-taxas obtido sobre a base de capital investido."
)

# Tabela EVA
data_eva = [
    ["Métrica Financeira (2025)", "ABEV3 (Ambev)", "ASAI3 (Assaí)", "AMER3 (Americanas)"],
    ["Retorno s/ Capital Investido (ROIC pós-tax)", "19,24%", "8,89%", "2,07%"],
    ["Custo Médio Ponderado de Capital (WACC)", "14,08%", "10,14%", "12,28%"],
    ["Spread Econômico (ROIC - WACC)", "+5,16%", "-1,25%", "-10,21%"],
    ["Base de Capital Investido Estimada", "R$ 80.380 mi", "R$ 26.640 mi", "R$ 10.520 mi"],
    ["Economic Value Added (EVA)", "+R$ 4.148 mi", "-R$ 333 mi", "-R$ 1.074 mi"]
]
create_table(doc, 6, 4, data_eva, col_widths)

add_body_paragraph(doc,
    "Os spreads econômicos revelam três dinâmicas completamente distintas em termos de geração de riqueza:"
)

add_body_paragraph(doc,
    "Spread Positivo (Ambev): A Ambev obteve spread econômico de +5,16%, resultante de um ROIC de 19,24% superando com folga o custo "
    "médio WACC de 14,08%. A rentabilidade de suas operações pagou todas as fontes de capital e gerou riqueza econômica líquida de "
    "R$ 4.148 milhões em 2025, justificando por que o mercado acionário a considera um ativo de refúgio e alta qualidade.",
    bold_prefix="* "
)

add_body_paragraph(doc,
    "Destruição por Custo de Capital (Assaí): O Assaí registrou spread negativo de -1,25%, decorrente de um ROIC de 8,89% inferior ao "
    "WACC de 10,14%. Embora o lucro operacional (EBIT) seja nominalmente positivo, a base de capital exigida para financiar as lojas "
    "convertidas do Extra tem um custo de oportunidade superior ao retorno das operações. Na prática, a empresa destruiu R$ 333 milhões "
    "de riqueza residual de seus sócios no ano.",
    bold_prefix="* "
)

add_body_paragraph(doc,
    "Destruição Condenável (Americanas): A Americanas destruiu R$ 1.074 milhões de valor acionário em 2025. O spread econômico foi "
    "de drásticos -10,21%, resultado de um ROIC ínfimo de 2,07% contra um WACC elevado de 12,28%, inflado pelo alto prêmio de risco "
    "exigido pelos credores da companhia em processo de recuperação judicial.",
    bold_prefix="* "
)

# ── PÁGINA 6: PERFORMANCE DE MERCADO E RISCO ACIONÁRIO ──
add_heading1(doc, "5. Performance de Mercado e Risco Acionário (2021-2025)")

add_body_paragraph(doc,
    "A saúde financeira refletida nos demonstrativos contábeis tem impacto direto na precificação das ações no mercado de capitais secundário. "
    "A tabela a seguir apresenta os indicadores estatísticos de performance e risco acionário baseados nas séries diárias completas ajustadas "
    "de ABEV3, ASAI3, AMER3 contra o benchmark do Ibovespa ao longo de 5 anos."
)

# Tabela Mercado
data_mercado = [
    ["Estatística Acionária (2021-2025)", "ABEV3", "AMER3", "ASAI3", "IBOV (Índice)"],
    ["Retorno Acumulado (5 anos)", "29,28%", "-99,94%", "-47,19%", "46,03%"],
    ["Retorno Anualizado (CAGR)", "5,27%", "-76,96%", "-11,99%", "7,87%"],
    ["Volatilidade Anualizada", "23,34%", "129,45%", "40,45%", "17,46%"],
    ["Índice Sharpe (Rf = 10,75%)", "-0,23", "-0,68", "-0,56", "-0,17"],
    ["Beta Sistemático (vs IBOV)", "0,6430", "1,5005", "1,0981", "1,0000"],
    ["Máximo Drawdown (Pior Queda)", "-36,18%", "-99,96%", "-75,55%", "-26,50%"]
]
col_widths_m = [Inches(3.4), Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.8)]
create_table(doc, 7, 5, data_mercado, col_widths_m)

add_body_paragraph(doc,
    "A curva de retorno acumulado normalizado exibe a dominância do benchmark: o Ibovespa registrou alta acumulada de 46,03% no "
    "período. A Ambev (ABEV3) operou com menor volatilidade (23,34% a.a.) e o menor Drawdown do grupo de ações (-36,18%), fechando o "
    "período com retorno positivo de 29,28%, consolidando-se como ativo de proteção patrimonial."
)

add_body_paragraph(doc,
    "O Assaí (ASAI3) registrou desvalorização acumulada de -47,19% no período, reflexo da aversão de risco diante do endividamento bruto "
    "elevado. Sua volatilidade anualizada de 40,45% foi mais que o dobro do Ibovespa, com queda máxima acumulada (Drawdown) de -75,55% de "
    "suas cotações."
)

add_body_paragraph(doc,
    "As Lojas Americanas (AMER3) sofreram uma ruína acionária quase integral, registrando perda de -99,94% de seu valor de mercado. "
    "Sua volatilidade anualizada extrema de 129,45% a.a. e a perda de capital acionário acentuada refletem o risco extremo de desastre "
    "administrativo-contábil, com Drawdown de -99,96%."
)

add_heading2(doc, "5.1. Matriz de Correlação de Retornos Diários")
add_body_paragraph(doc,
    "A correlação estatística cruzada de retornos diários revela o grau de movimentação conjunta das ações contra o Ibovespa."
)

# Tabela Correlação
data_corr = [
    ["Matriz de Correlação", "ABEV3", "AMER3", "ASAI3", "IBOV"],
    ["ABEV3", "1,0000", "0,0878", "0,2511", "0,4809"],
    ["AMER3", "0,0878", "1,0000", "0,1331", "0,2023"],
    ["ASAI3", "0,2511", "0,1331", "1,0000", "0,4738"],
    ["IBOV", "0,4809", "0,2023", "0,4738", "1,0000"]
]
create_table(doc, 5, 5, data_corr, col_widths_m)

add_body_paragraph(doc,
    "A baixa correlação cruzada diária entre as três ações (por exemplo, 0,08 entre ABEV3 e AMER3) sugere que os preços das ações "
    "reagiram de forma individualizada a notícias corporativas microeconômicas específicas ao invés de oscilarem de forma sistemática. "
    "Para alocação em portfólios diversificados, a Ambev desponta como excelente diversificador estratégico devido ao seu Beta sistemático "
    "baixo de 0,6430 e correlação moderada com o Ibovespa (0,48)."
)

# ── PÁGINA 7: PROJEÇÕES E SIMULAÇÃO MONTE CARLO ──
add_heading1(doc, "6. Projeção de Risco Futuro via Simulação de Monte Carlo (2026)")

add_body_paragraph(doc,
    "Para estimar a dispersão e o risco futuro de mercado para o ano de 2026, implementou-se uma modelagem de projeção estocástica baseada "
    "no Movimento Browniano Geométrico (MBG). O modelo é calibrado utilizando a tendência (drift) histórica e a volatilidade (sigma) "
    "calculadas a partir das séries completas de 5 anos."
)

add_body_paragraph(doc,
    "A equação estocástica do MBG assume que as variações percentuais nos preços das ações seguem uma componente determinística de tendência "
    "e um choque aleatório representado pelo processo de Wiener (distribuição normal de média zero e desvio-padrão unitário). Foram simulados "
    "10.000 caminhos alternativos de preços para 252 dias úteis de mercado de capitais do ano de 2026."
)

# Tabela Monte Carlo
data_mc = [
    ["Estatística Monte Carlo (2026)", "ABEV3 (Ambev)", "AMER3 (Americanas)", "ASAI3 (Assaí)"],
    ["Último Preço Real (Dez/2025)", "R$ 13,86", "R$ 5,13", "R$ 7,20"],
    ["Preço Mediano Projetado", "R$ 14,26 (+2,88%)", "R$ 1,11 (-78,38%)", "R$ 6,30 (-12,54%)"],
    ["Percentil 5% (Pior Cenário)", "R$ 9,67", "R$ 0,14", "R$ 3,23"],
    ["Percentil 95% (Melhor Cenário)", "R$ 20,99", "R$ 9,05", "R$ 12,26"],
    ["Value at Risk (VaR 95% Anual)", "-30,27%", "-97,26%", "-55,09%"],
    ["Probabilidade de Perda Anual", "45,1%", "88,6%", "63,0%"]
]
create_table(doc, 7, 4, data_mc, col_widths)

add_body_paragraph(doc,
    "Os resultados das projeções estocásticas confirmam e quantificam de forma matemática a saúde contábil de cada empresa:"
)

add_body_paragraph(doc,
    "ABEV3 (Risco Controlado): O funil estocástico da Ambev é estreito e centrado. Partindo de R$ 13,86, a mediana projeta R$ 14,26 "
    "em dezembro de 2026. O pior cenário (P5) estima R$ 9,67. O VaR de -30,27% expressa a menor perda esperada e a probabilidade de perda "
    "é de apenas 45,1%, demonstrando segurança relativa para o investidor.",
    bold_prefix="* "
)

add_body_paragraph(doc,
    "ASAI3 (Risco Intermediário): O funil do Assaí reflete volatilidade moderada-alta. O preço mediano esperado é de R$ 6,30 (-12,54%). "
    "O VaR anualizado de -55,09% sinaliza que no pior cenário a ação cairia para R$ 3,23. A probabilidade de perda no ano é de 63,0%, refletindo "
    "o peso continuado dos juros de financiamento no preço do ativo.",
    bold_prefix="* "
)

add_body_paragraph(doc,
    "AMER3 (Risco Crítico): O funil de Americanas é extremamente disperso e assimétrico, com tendência majoritária de convergência para "
    "próximo de zero. A mediana projetada é de R$ 1,11 (desvalorização de -78,38%). A probabilidade de perda em 2026 é de 88,6%, com VaR de "
    "-97,26% (preço caindo para R$ 0,14 no pior cenário). O papel opera sob caráter especulativo extremo com alta probabilidade de colapso acionário.",
    bold_prefix="* "
)

# ── PÁGINA 8: CONCLUSÕES E RECOMENDAÇÕES ──
add_heading1(doc, "7. Conclusões Integradas e Recomendação Prática")

add_body_paragraph(doc,
    "A integração analítica contábil, de criação de valor residual e estocástica de mercado construída neste trabalho permite chegar a "
    "quatro conclusões fundamentais para tomadores de decisão financeira e gestores de investimento:"
)

add_body_paragraph(doc,
    "A saúde dos demonstrativos contábeis determina o comportamento de longo prazo na Bolsa de Valores. A resiliência da Ambev "
    "decorre de seu EVA recorrentemente positivo (+R$ 4.148 milhões) e de sua estrutura de capital livre de dívidas, blindando as "
    "ações contra pressões macroeconômicas de taxas de juros elevadas.",
    bold_prefix="1. "
)

add_body_paragraph(doc,
    "A expansão corporativa financiada por alavancagem financeira extrema em ciclos de juros altos representa um elevado risco "
    "de mercado. O Assaí demonstra que o faturamento elevado (R$ 71,5 bilhões) é insuficiente se a rentabilidade operacional "
    "(ROIC) for menor que o custo médio de capital ponderado (WACC), gerando spread negativo e destruição econômica (EVA negativo).",
    bold_prefix="2. "
)

add_body_paragraph(doc,
    "A fragilidade administrativa e as falhas profundas de governança corporativa produzem perdas binárias de capital permanente. "
    "O colapso da Americanas (-99,94% de retornos) materializa o risco extremo de crédito e destruição patrimonial irreversível.",
    bold_prefix="3. "
)

add_body_paragraph(doc,
    "A modelagem estocástica por simulação de Monte Carlo serve como rica ferramenta complementar de gerenciamento de risco de portfólio. "
    "Ao calibrar drift e volatilidade de longo prazo, o modelo estocástico prevê matematicamente a dispersão das cotações futuras, "
    "auxiliando na estimativa de margem de segurança e Value at Risk (VaR).",
    bold_prefix="4. "
)

add_body_paragraph(doc,
    "Como diretriz prática de investimento, recomenda-se a priorização na alocação de recursos em corporações com spreads econômicos "
    "positivos e baixa alavancagem financeira, exigindo margem de segurança baseada no VaR anualizado para mitigar os riscos em momentos "
    "de juros de dois dígitos no Brasil."
)

# ── PÁGINA 9: REFERÊNCIAS BIBLIOGRÁFICAS ──
add_heading1(doc, "8. Referências Bibliográficas")
add_empty_line(doc)

add_reference_abnt(doc,
    "AMERICANAS S.A.", "Central de Resultados.",
    " Rio de Janeiro: Lojas Americanas Relações com Investidores, 2025.",
    "https://ri.americanas.io/informacoes-aos-investidores/central-de-resultados/"
)

add_reference_abnt(doc,
    "AMBEV S.A.", "Divulgação de Resultados.",
    " São Paulo: Ambev Relações com Investidores, 2025.",
    "https://ri.ambev.com.br/relatorios-publicacoes/divulgacao-de-resultados/"
)

add_reference_abnt(doc,
    "ASSAÍ ATACADISTA.", "Resultados Trimestrais.",
    " São Paulo: Assaí Relações com Investidores, 2025.",
    "https://ri.assai.com.br/informacoes-financeiras/resultados-trimestrais/"
)

add_reference_abnt(doc,
    "ASSAF NETO, Alexandre.", "Finanças Corporativas e Valor.",
    " 8. ed. São Paulo: Atlas, 2020. ISBN 978-85-97-02637-5."
)

add_reference_abnt(doc,
    "B3 — BRASIL, BOLSA, BALCÃO.", "Histórico de cotações de ativos listados.",
    " São Paulo: B3 S.A., 2025.",
    "https://www.b3.com.br"
)

add_reference_abnt(doc,
    "BANCO CENTRAL DO BRASIL.", "Taxa Selic — Dados Diários e Históricos.",
    " Brasília: BCB, 2025.",
    "https://www.bcb.gov.br/estatisticas/grafico/graficoestatistica/taxaselic"
)

add_reference_abnt(doc,
    "BRASIL. Comissão de Valores Mobiliários — CVM.", "Dados cadastrais e demonstrativos financeiros de empresas abertas.",
    " Rio de Janeiro: CVM, 2025.",
    "https://dados.cvm.gov.br"
)

add_reference_abnt(doc,
    "BREALEY, Richard A.; MYERS, Stewart C.; ALLEN, Franklin.", "Principles of Corporate Finance.",
    " 14. ed. New York: McGraw-Hill Education, 2022. ISBN 978-1-260-56555-6."
)

add_reference_abnt(doc,
    "DAMODARAN, Aswath.", "Investment Valuation: Tools and Techniques for Determining the Value of Any Asset.",
    " 3. ed. Hoboken: John Wiley & Sons, 2012. ISBN 978-1-118-01152-2."
)

add_reference_abnt(doc,
    "FLEURIET, Michel; KEHDY, Ricardo; BLANC, Georges.", "O Modelo Fleuriet: A Dinâmica Financeira das Empresas Brasileiras.",
    " 7. ed. Rio de Janeiro: Elsevier, 2003. ISBN 978-85-352-1237-4."
)

add_reference_abnt(doc,
    "GLASSERMAN, Paul.", "Monte Carlo Methods in Financial Engineering.",
    " New York: Springer, 2004. (Applications of Mathematics, v. 53). ISBN 978-0-387-00451-8."
)

add_reference_abnt(doc,
    "HULL, John C.", "Options, Futures, and Other Derivatives.",
    " 11. ed. Harlow: Pearson Education, 2022. ISBN 978-0-136-93999-5."
)

add_reference_abnt(doc,
    "JORION, Philippe.", "Value at Risk: The New Benchmark for Managing Financial Risk.",
    " 3. ed. New York: McGraw-Hill, 2007. ISBN 978-0-071-46495-6."
)

add_reference_abnt(doc,
    "LABORATÓRIO DE FINANÇAS.", "API de Preços Corrigidos — v2.",
    " Plataforma acadêmica para extração de séries históricas de cotações ajustadas e indicadores contábeis, 2026.",
    "https://laboratoriodefinancas.com"
)

add_reference_abnt(doc,
    "MARKOWITZ, Harry M.", "Portfolio Selection.",
    " The Journal of Finance, v. 7, n. 1, p. 77-91, mar. 1952. DOI: 10.2307/2975974."
)

add_reference_abnt(doc,
    "SHARPE, William F.", "Capital Asset Prices: A Theory of Market Equilibrium under Conditions of Risk.",
    " The Journal of Finance, v. 19, n. 3, p. 425-442, set. 1964. DOI: 10.2307/2977928."
)

add_reference_abnt(doc,
    "STEWART, G. Bennett III.", "The Quest for Value: A Guide for Senior Managers.",
    " New York: Harper Business, 1991. ISBN 978-0-887-30418-1."
)

add_empty_line(doc)

# Rodapé final centralizado
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = p_end.add_run("Norma Técnica de Referência: ABNT NBR 6023:2018 — Informação e documentação — Referências.")
run_end.italic = True
run_end.font.name = FONT_NAME
run_end.font.size = Pt(9.5)
run_end.font.color.rgb = GRAY_TEXT

# Salva o documento
output_path = r"c:\andd\ap2_analise_de_dados\1_Relatorio_Escrito_Completo.docx"
doc.save(output_path)
print(f"Trabalho completo salvo com sucesso em: {output_path}")
