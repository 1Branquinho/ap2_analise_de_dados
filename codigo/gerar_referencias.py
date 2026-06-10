import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ════════════════════════════════════════════════════════════════════
# Abre o documento original como template para herdar estilos
# ════════════════════════════════════════════════════════════════════
template = Document(r'c:\andd\ap2_analise_de_dados\4_Referencias_Bibliograficas.docx')
doc = Document()

# Copia as dimensões de página e margens do template
for sec_src, sec_dst in zip(template.sections, doc.sections):
    sec_dst.page_width = sec_src.page_width
    sec_dst.page_height = sec_src.page_height
    sec_dst.left_margin = sec_src.left_margin
    sec_dst.right_margin = sec_src.right_margin
    sec_dst.top_margin = sec_src.top_margin
    sec_dst.bottom_margin = sec_src.bottom_margin

# ════════════════════════════════════════════════════════════════════
# Funções auxiliares para replicar o estilo do documento original
# ════════════════════════════════════════════════════════════════════

FONT_NAME = "Arial"
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)   # #1F3864
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LINK_BLUE = RGBColor(0x11, 0x55, 0xCC)   # #1155CC
GRAY = RGBColor(0x88, 0x88, 0x88)

def add_heading1(doc, text):
    """Título principal do documento — Heading 1 com fundo azul escuro."""
    p = doc.add_paragraph()
    # Aplica shading (fundo azul escuro)
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:color="1F3864"/>')
    pPr.append(shd)
    # Espaçamento
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
    pPr.append(spacing)
    # Indentação
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200" w:right="200"/>')
    pPr.append(ind)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.color.rgb = WHITE
    return p

def add_heading2(doc, text):
    """Subtítulo de seção — Heading 2 com fundo azul escuro."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:color="1F3864"/>')
    pPr.append(shd)
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="400"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200" w:right="200"/>')
    pPr.append(ind)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    return p

def add_intro_text(doc, text):
    """Parágrafo introdutório em itálico."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
    pPr.append(spacing)
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    return p

def add_reference(doc, author, title, rest, url=None):
    """Entrada de referência bibliográfica no padrão ABNT."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="720"/>')
    pPr.append(ind)
    
    # Autor (negrito)
    r_author = p.add_run(author)
    r_author.bold = True
    r_author.font.name = FONT_NAME
    r_author.font.size = Pt(12)
    
    # Espaço
    r_space = p.add_run(" ")
    r_space.font.name = FONT_NAME
    r_space.font.size = Pt(12)
    
    # Título (negrito)
    r_title = p.add_run(title)
    r_title.bold = True
    r_title.font.name = FONT_NAME
    r_title.font.size = Pt(12)
    
    # Texto restante (normal)
    r_rest = p.add_run(rest)
    r_rest.font.name = FONT_NAME
    r_rest.font.size = Pt(12)
    
    # URL como hyperlink (se fornecida)
    if url:
        add_hyperlink(p, url)
        r_after = p.add_run(". Acesso em: 10 jun. 2026.")
        r_after.font.name = FONT_NAME
        r_after.font.size = Pt(12)
    
    return p

def add_hyperlink(paragraph, url, text=None):
    """Adiciona um hyperlink ao parágrafo."""
    if text is None:
        text = url
    
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial" w:eastAsia="Arial"/>'
        f'    <w:color w:val="1155CC"/>'
        f'    <w:sz w:val="24"/>'
        f'    <w:szCs w:val="24"/>'
        f'    <w:u w:val="single"/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def add_nota_title(doc, text):
    """Título estilo 'Nota Metodológica' — azul escuro sem fundo."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="400"/>')
    pPr.append(spacing)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_BLUE
    return p

def add_body_text(doc, text, alignment=None):
    """Parágrafo de corpo justificado."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
    pPr.append(spacing)
    
    if alignment == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    return p

def add_footer_note(doc, text):
    """Nota de rodapé centralizada em itálico cinza."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(text)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p

def add_empty_line(doc):
    """Linha em branco."""
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    return p

# ════════════════════════════════════════════════════════════════════
# Remove o parágrafo padrão vazio que o python-docx cria
# ════════════════════════════════════════════════════════════════════
if doc.paragraphs:
    doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

# ════════════════════════════════════════════════════════════════════
# CONSTRUÇÃO DO DOCUMENTO
# ════════════════════════════════════════════════════════════════════

print("Gerando documento de referências bibliográficas...")

# ── TÍTULO PRINCIPAL ──
add_heading1(doc, "Análise de Desempenho do Setor de Varejo e Consumo Brasileiro")

add_empty_line(doc)

# ── INTRODUÇÃO ──
add_intro_text(doc,
    "As referências listadas a seguir fundamentam a análise financeira, setorial, de mercado e de risco realizada neste trabalho, "
    "abrangendo dados contábeis (DRE, Balanço Patrimonial, Fleuriet), indicadores de performance acionária (retornos, volatilidade, Beta, Sharpe), "
    "criação de valor econômico (EVA/ROIC/WACC) e projeções estocásticas via simulação de Monte Carlo. "
    "As fontes estão organizadas conforme suas categorias de origem."
)

add_empty_line(doc)

# ═══════════════════════════════════════════════════════════════
# SEÇÃO 1: Fontes Primárias — Relações com Investidores
# ═══════════════════════════════════════════════════════════════
add_heading2(doc, "1. Fontes Primárias — Relações com Investidores")
add_empty_line(doc)

add_reference(doc,
    "AMERICANAS S.A.", "Central de Resultados.",
    " Disponível em: ",
    "https://ri.americanas.io/informacoes-aos-investidores/central-de-resultados/"
)

add_reference(doc,
    "AMBEV S.A.", "Divulgação de Resultados.",
    " Disponível em: ",
    "https://ri.ambev.com.br/relatorios-publicacoes/divulgacao-de-resultados/"
)

add_reference(doc,
    "ASSAÍ ATACADISTA.", "Resultados Trimestrais.",
    " Disponível em: ",
    "https://ri.assai.com.br/informacoes-financeiras/resultados-trimestrais/"
)

# ═══════════════════════════════════════════════════════════════
# SEÇÃO 2: Dados Regulatórios e de Mercado
# ═══════════════════════════════════════════════════════════════
add_heading2(doc, "2. Dados Regulatórios e de Mercado")
add_empty_line(doc)

add_reference(doc,
    "BRASIL. Comissão de Valores Mobiliários — CVM.", "Dados de empresas abertas.",
    " Disponível em: ",
    "https://dados.cvm.gov.br"
)

add_reference(doc,
    "B3 — BRASIL, BOLSA, BALCÃO.", "Histórico de cotações.",
    " Disponível em: ",
    "https://www.b3.com.br"
)

add_reference(doc,
    "LABORATÓRIO DE FINANÇAS.", "API de Preços Corrigidos — v2.",
    " Plataforma acadêmica para extração de séries históricas de cotações ajustadas e indicadores contábeis. Disponível em: ",
    "https://laboratoriodefinancas.com"
)

add_reference(doc,
    "BANCO CENTRAL DO BRASIL.", "Taxa Selic — Dados Diários.",
    " Disponível em: ",
    "https://www.bcb.gov.br/estatisticas/grafico/graficoestatistica/taxaselic"
)

# ═══════════════════════════════════════════════════════════════
# SEÇÃO 3: Referencial Teórico — Finanças Corporativas
# ═══════════════════════════════════════════════════════════════
add_heading2(doc, "3. Referencial Teórico — Finanças Corporativas e Valuation")
add_empty_line(doc)

add_reference(doc,
    "DAMODARAN, Aswath.", "Investment Valuation: Tools and Techniques for Determining the Value of Any Asset.",
    " 3. ed. Hoboken: John Wiley & Sons, 2012. ISBN 978-1-118-01152-2.",
)

add_reference(doc,
    "ASSAF NETO, Alexandre.", "Finanças Corporativas e Valor.",
    " 8. ed. São Paulo: Atlas, 2020. ISBN 978-85-97-02637-5.",
)

add_reference(doc,
    "BREALEY, Richard A.; MYERS, Stewart C.; ALLEN, Franklin.", "Principles of Corporate Finance.",
    " 14. ed. New York: McGraw-Hill Education, 2022. ISBN 978-1-260-56555-6.",
)

add_reference(doc,
    "STEWART, G. Bennett III.", "The Quest for Value: A Guide for Senior Managers.",
    " New York: Harper Business, 1991. ISBN 978-0-887-30418-1.",
)

add_reference(doc,
    "FLEURIET, Michel; KEHDY, Ricardo; BLANC, Georges.", "O Modelo Fleuriet: A Dinâmica Financeira das Empresas Brasileiras.",
    " 7. ed. Rio de Janeiro: Elsevier, 2003. ISBN 978-85-352-1237-4.",
)

# ═══════════════════════════════════════════════════════════════
# SEÇÃO 4: Referencial Teórico — Análise de Risco e Mercado
# ═══════════════════════════════════════════════════════════════
add_heading2(doc, "4. Referencial Teórico — Análise de Risco e Mercado de Capitais")
add_empty_line(doc)

add_reference(doc,
    "HULL, John C.", "Options, Futures, and Other Derivatives.",
    " 11. ed. Harlow: Pearson Education, 2022. ISBN 978-0-136-93999-5.",
)

add_reference(doc,
    "SHARPE, William F.", "Capital Asset Prices: A Theory of Market Equilibrium under Conditions of Risk.",
    " The Journal of Finance, v. 19, n. 3, p. 425-442, set. 1964. DOI: 10.2307/2977928.",
)

add_reference(doc,
    "MARKOWITZ, Harry M.", "Portfolio Selection.",
    " The Journal of Finance, v. 7, n. 1, p. 77-91, mar. 1952. DOI: 10.2307/2975974.",
)

add_reference(doc,
    "GLASSERMAN, Paul.", "Monte Carlo Methods in Financial Engineering.",
    " New York: Springer, 2004. (Applications of Mathematics, v. 53). ISBN 978-0-387-00451-8.",
)

add_reference(doc,
    "JORION, Philippe.", "Value at Risk: The New Benchmark for Managing Financial Risk.",
    " 3. ed. New York: McGraw-Hill, 2007. ISBN 978-0-071-46495-6.",
)

# ═══════════════════════════════════════════════════════════════
# SEÇÃO 5: Fontes Audiovisuais
# ═══════════════════════════════════════════════════════════════
add_heading2(doc, "5. Fontes Audiovisuais")
add_empty_line(doc)

# Vídeo 1
p = doc.add_paragraph()
pPr = p._element.get_or_add_pPr()
spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
pPr.append(spacing)
ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="720"/>')
pPr.append(ind)

r1 = p.add_run("ECONOMIA E VAREJO: COMO 2025 FECHA E AS EXPECTATIVAS PARA 2026")
r1.bold = True; r1.font.name = FONT_NAME; r1.font.size = Pt(12)
r2 = p.add_run(" | Gabriel Couto, Santander. ")
r2.font.name = FONT_NAME; r2.font.size = Pt(12)
r3 = p.add_run("[Vídeo]. ")
r3.font.name = FONT_NAME; r3.font.size = Pt(12)
r4 = p.add_run("YouTube,")
r4.bold = True; r4.font.name = FONT_NAME; r4.font.size = Pt(12)
r5 = p.add_run(" 2025. Disponível em: ")
r5.font.name = FONT_NAME; r5.font.size = Pt(12)
add_hyperlink(p, "https://www.youtube.com/watch?v=dQw4w9WgXcQ")
r6 = p.add_run(". Acesso em: 10 jun. 2026.")
r6.font.name = FONT_NAME; r6.font.size = Pt(12)

# Vídeo 2
p = doc.add_paragraph()
pPr = p._element.get_or_add_pPr()
spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
pPr.append(spacing)
ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="720"/>')
pPr.append(ind)

r1 = p.add_run("SETOR VAREJISTA NO BRASIL: O QUE ESTÁ ACONTECENDO?")
r1.bold = True; r1.font.name = FONT_NAME; r1.font.size = Pt(12)
r2 = p.add_run(" [Vídeo]. ")
r2.font.name = FONT_NAME; r2.font.size = Pt(12)
r3 = p.add_run("YouTube,")
r3.bold = True; r3.font.name = FONT_NAME; r3.font.size = Pt(12)
r4 = p.add_run(" 2025. Disponível em: ")
r4.font.name = FONT_NAME; r4.font.size = Pt(12)
add_hyperlink(p, "https://youtu.be/5jyqnw0FlQw")
r5 = p.add_run(". Acesso em: 10 jun. 2026.")
r5.font.name = FONT_NAME; r5.font.size = Pt(12)

# ═══════════════════════════════════════════════════════════════
# SEÇÃO 6: Ferramentas Computacionais
# ═══════════════════════════════════════════════════════════════
add_heading2(doc, "6. Ferramentas Computacionais e Bibliotecas")
add_empty_line(doc)

add_reference(doc,
    "PYTHON SOFTWARE FOUNDATION.", "Python Language Reference, version 3.14.",
    " Disponível em: ",
    "https://www.python.org"
)

add_reference(doc,
    "THE PANDAS DEVELOPMENT TEAM.", "pandas: powerful Python data analysis toolkit.",
    " Disponível em: ",
    "https://pandas.pydata.org"
)

add_reference(doc,
    "HUNTER, John D.", "Matplotlib: A 2D graphics environment.",
    " Computing in Science & Engineering, v. 9, n. 3, p. 90-95, 2007. DOI: 10.1109/MCSE.2007.55. Disponível em: ",
    "https://matplotlib.org"
)

add_reference(doc,
    "HARRIS, Charles R. et al.", "Array programming with NumPy.",
    " Nature, v. 585, p. 357-362, 2020. DOI: 10.1038/s41586-020-2649-2. Disponível em: ",
    "https://numpy.org"
)

add_reference(doc,
    "PYTHON-PPTX.", "python-pptx: Generate and manipulate Open XML PowerPoint.",
    " Disponível em: ",
    "https://python-pptx.readthedocs.io"
)

add_empty_line(doc)

# ═══════════════════════════════════════════════════════════════
# NOTA METODOLÓGICA
# ═══════════════════════════════════════════════════════════════
add_nota_title(doc, "Nota Metodológica")

add_body_text(doc,
    "Os dados financeiros extraídos das fontes primárias foram processados utilizando a linguagem Python e a biblioteca Pandas. "
    "As séries históricas de cotações ajustadas (2021-2025) foram obtidas via API do Laboratório de Finanças. "
    "Os indicadores de balanço (DRE, BP, EBITDA, EBIT, Margem Líquida, Liquidez, Fleuriet, Endividamento) foram calculados a partir dos demonstrativos financeiros padronizados (DFP). "
    "O cálculo do EVA (Economic Value Added) seguiu a metodologia de Stern Stewart (1991), com WACC estimado pelo modelo CAPM de Sharpe (1964) e dados de risco-país e prêmio de mercado de Damodaran. "
    "A análise de mercado incluiu retorno acumulado normalizado, volatilidade anualizada, índice de Sharpe, Beta sistemático (regressão OLS contra IBOV) e máximo drawdown. "
    "A simulação de Monte Carlo utilizou o Movimento Browniano Geométrico (GBM) com 10.000 simulações de 252 dias úteis, "
    "calculando percentis, Value at Risk (VaR) a 95% de confiança e probabilidade de perda, "
    "conforme metodologia descrita em Glasserman (2004) e Jorion (2007). "
    "A apresentação de slides foi gerada programaticamente com python-pptx em formato widescreen 16:9.",
    alignment="justify"
)

add_footer_note(doc, "Norma de referência: ABNT NBR 6023:2018 — Informação e documentação — Referências — Elaboração.")

# ════════════════════════════════════════════════════════════════════
# SALVA O DOCUMENTO
# ════════════════════════════════════════════════════════════════════
output_path = r"c:\andd\ap2_analise_de_dados\4_Referencias_Bibliograficas.docx"
doc.save(output_path)
print(f"Documento salvo com sucesso: {output_path}")
