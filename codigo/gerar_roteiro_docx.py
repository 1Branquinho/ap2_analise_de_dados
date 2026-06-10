import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ════════════════════════════════════════════════════════════════════
# Inicializa o documento e copia as configurações de página
# ════════════════════════════════════════════════════════════════════
template = Document(r'c:\andd\ap2_analise_de_dados\4_Referencias_Bibliograficas.docx')
doc = Document()

for sec_src, sec_dst in zip(template.sections, doc.sections):
    sec_dst.page_width = sec_src.page_width
    sec_dst.page_height = sec_src.page_height
    sec_dst.left_margin = sec_src.left_margin
    sec_dst.right_margin = sec_src.right_margin
    sec_dst.top_margin = sec_src.top_margin
    sec_dst.bottom_margin = sec_src.bottom_margin

# ════════════════════════════════════════════════════════════════════
# Configuração de Estilos e Paleta de Cores
# ════════════════════════════════════════════════════════════════════
FONT_NAME = "Arial"
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)   # #1F3864
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_TEXT = RGBColor(0x59, 0x59, 0x59)   # #595959
SPEECH_COLOR = RGBColor(0x26, 0x26, 0x26) # #262626
LINK_BLUE = RGBColor(0x11, 0x55, 0xCC)

def add_title(doc, text):
    """Título principal do documento com fundo azul escuro."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:color="1F3864"/>')
    pPr.append(shd)
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="240" w:before="0"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200" w:right="200"/>')
    pPr.append(ind)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.color.rgb = WHITE
    return p

def add_section_header(doc, text):
    """Título de Seção (Heading 2) com fundo azul escuro."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:color="1F3864"/>')
    pPr.append(shd)
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="160" w:before="360"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200" w:right="200"/>')
    pPr.append(ind)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    return p

def add_subsection_header(doc, text):
    """Subtítulo de seção em azul escuro sem fundo."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="120" w:before="240"/>')
    pPr.append(spacing)
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    return p

def add_meta_info(doc, slide_info, objective):
    """Informações de metadados do orador (Slide, Objetivo) em bloco com recuo."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="100" w:before="0"/>')
    pPr.append(spacing)
    
    # Borda esquerda cinza para destacar o metadado
    pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="A6A6A6"/></w:pBdr>')
    pPr.append(pbdr)
    
    r_slide = p.add_run("📊 ID do Slide: ")
    r_slide.bold = True
    r_slide.font.name = FONT_NAME
    r_slide.font.size = Pt(9.5)
    r_slide.font.color.rgb = GRAY_TEXT
    
    r_slide_val = p.add_run(f"{slide_info}\n")
    r_slide_val.font.name = FONT_NAME
    r_slide_val.font.size = Pt(9.5)
    r_slide_val.font.color.rgb = GRAY_TEXT
    
    r_obj = p.add_run("🎯 Objetivo: ")
    r_obj.bold = True
    r_obj.font.name = FONT_NAME
    r_obj.font.size = Pt(9.5)
    r_obj.font.color.rgb = GRAY_TEXT
    
    r_obj_val = p.add_run(objective)
    r_obj_val.font.name = FONT_NAME
    r_obj_val.font.size = Pt(9.5)
    r_obj_val.font.color.rgb = GRAY_TEXT
    return p

def add_speech(doc, text):
    """Adiciona o parágrafo de fala do orador (justificado e formatado)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="160" w:before="0"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:right="360"/>')
    pPr.append(ind)
    
    # Aspas duplas bonitas abrindo a fala
    r_quote_open = p.add_run("“ ")
    r_quote_open.bold = True
    r_quote_open.font.name = FONT_NAME
    r_quote_open.font.size = Pt(13)
    r_quote_open.font.color.rgb = DARK_BLUE
    
    r_text = p.add_run(text)
    r_text.font.name = FONT_NAME
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = SPEECH_COLOR
    
    r_quote_close = p.add_run(" ”")
    r_quote_close.bold = True
    r_quote_close.font.name = FONT_NAME
    r_quote_close.font.size = Pt(13)
    r_quote_close.font.color.rgb = DARK_BLUE
    return p

def add_tip(doc, text):
    """Dica prática de palco ou dinâmica em itálico cinza."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="200" w:before="0"/>')
    pPr.append(spacing)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360"/>')
    pPr.append(ind)
    
    run_label = p.add_run("💡 Dica de Apresentação: ")
    run_label.bold = True
    run_label.font.name = FONT_NAME
    run_label.font.size = Pt(9.5)
    run_label.font.color.rgb = GRAY_TEXT
    
    run_text = p.add_run(text)
    run_text.italic = True
    run_text.font.name = FONT_NAME
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = GRAY_TEXT
    return p

def add_intro_text(doc, text):
    """Parágrafo introdutório em itálico."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:after="240" w:before="0"/>')
    pPr.append(spacing)
    
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11.5)
    run.italic = True
    return p

def add_empty_line(doc):
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    return p

# Remove o parágrafo vazio padrão inicial
if doc.paragraphs:
    doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

# ════════════════════════════════════════════════════════════════════
# Geração do Conteúdo
# ════════════════════════════════════════════════════════════════════

print("Construindo o documento Word do Roteiro...")

# Título Principal
add_title(doc, "Roteiro de Falas e Guia Explicativo da Apresentação")
add_empty_line(doc)

add_intro_text(doc,
    "Este material serve como guia de apoio para o orador durante a apresentação do projeto. "
    "Conforme solicitado, a estrutura foi dividida em Falas Gerais (introdução, contexto, conceitos transversais e conclusões) "
    "e Blocos Temáticos específicos por empresa (Ambev, Assaí e Americanas). Isso permite que cada corporação seja discutida de forma consolidada e fluida."
)

# ────────────────────────────────────────────────────────────────────
# SEÇÃO 1: FALAS GERAIS (INDEPENDENTE / TRANSVERSAL)
# ────────────────────────────────────────────────────────────────────
add_section_header(doc, "Seção 1: Falas Gerais e Transversais")

# 1. Capa
add_subsection_header(doc, "1. Abertura do Trabalho")
add_meta_info(doc, "Slide 1 (Capa)", "Apresentar o tema do estudo, a equipe e introduzir a abordagem integrada contábil-estocástica.")
add_speech(doc,
    "Olá a todos. Vamos dar início à nossa apresentação executiva que aborda um estudo profundo de criação de valor "
    "econômico e performance de mercado no varejo e consumo brasileiro. O nosso objetivo é analisar e comparar o "
    "desempenho operacional, contábil e acionário de três gigantes do nosso mercado: a Ambev (ABEV3), o Assaí Atacadista "
    "(ASAI3) e as Lojas Americanas (AMER3), ao longo do período de 5 anos compreendido entre 2021 e 2025. "
    "Mais do que uma simples análise de indicadores tradicionais, o diferencial do nosso projeto está na integração multicritério. "
    "Nós cruzamos os dados contábeis puros com métricas de criação de valor residual, como o EVA (Economic Value Added), "
    "confrontamos esses resultados com a performance real das ações na Bolsa frente ao Ibovespa e, finalmente, rodamos "
    "uma projeção estocástica baseada em 10.000 simulações de Monte Carlo para mensurar matematicamente o risco estatístico "
    "de cada ativo para o ano de 2026. Vamos entender como a saúde interna das demonstrações financeiras molda as projeções de risco de mercado na vida real."
)
add_tip(doc, "Fale num tom pausado e profissional. Use o Slide 1 projetado como fundo enquanto contextualiza a relevância do estudo.")

# 2. Cenário Macro
add_subsection_header(doc, "2. Cenário Macroeconômico Brasileiro (2021-2025)")
add_meta_info(doc, "Slide 2 (Introdução e Dinâmica Macroeconômica)", "Criar o pano de fundo de juros e inflação que justifica a saúde financeira de cada empresa.")
add_speech(doc,
    "Antes de entrarmos nos números de cada companhia, precisamos entender sob qual ambiente macroeconômico elas operaram de 2021 a 2025. "
    "O mercado brasileiro foi marcado por três grandes forças. Primeiro: Juros Elevados. A taxa básica Selic permaneceu "
    "em dois dígitos na maior parte do tempo, encerrando o período na média de 10,75% ao ano. Isso encarece diretamente "
    "o custo das dívidas e aumenta a taxa livre de risco exigida pelos acionistas. Segundo: Pressão Inflacionária. "
    "O aumento de custos logísticos e de insumos industriais pressionou a renda real e exigiu das empresas varejistas um "
    "rigor operacional absurdo para não perder volume. Terceiro: Aversão ao Risco. Diante de tanta volatilidade, o mercado passou "
    "a punir severamente empresas com alto endividamento e a premiar aquelas que geram caixa próprio constante. "
    "Sob essa tempestade contínua, escolhemos três empresas com perfis diametralmente opostos de capital. A Ambev representa o "
    "consumo não cíclico e defensivo. O Assaí atua no varejo alimentar de alto giro e baixas margens, que expandiu via dívidas. "
    "E a Americanas representa o varejo discricionário exposto à concorrência digital implacável e que sofreu uma profunda crise de governança."
)
add_tip(doc, "Destaque as palavras 'Juros Elevados' e 'Aversão ao Risco' para preparar o público para a análise de dívida das empresas.")

# 9. Tabela Contábil Consolidada
add_subsection_header(doc, "3. Painel Contábil Consolidado (2025)")
add_meta_info(doc, "Slide 9 (Painel Comparativo Contábil)", "Demonstrar a disparidade contábil direta entre as três empresas e introduzir os números da tabela.")
add_speech(doc,
    "Neste momento da nossa análise, trazemos um painel contábil consolidado das três companhias referente ao ano de 2025. "
    "Essa tabela unifica os dados extraídos das Demonstrações Financeiras Padronizadas (DFPs) e da CVM. "
    "O que chama a atenção imediatamente aqui é o abismo em termos de margens e endividamento. Enquanto a Ambev lidera o faturamento "
    "nominal com 82,5 bilhões de reais de Receita Líquida e sustenta uma margem de Lucro Operacional (EBIT) robusta de 26,54%, o Assaí "
    "possui uma receita gigantesca de 71,5 bilhões, mas trabalha com a margem operacional reduzida típica do atacarejo, de 5,07%. "
    "Na ponta mais fraca, a Americanas reporta uma receita encolhida de 12,3 bilhões de reais com margem operacional de 2,69% e prejuízo "
    "líquido de 271 milhões. Mas o grande divisor de águas reside na linha de Dívida Líquida. Observem que a Ambev possui caixa líquido positivo "
    "de 8,4 bilhões de reais. O Assaí, por sua vez, carrega dívida líquida imensa de 21,3 bilhões de reais, e a Americanas, 5,7 bilhões."
)
add_tip(doc, "Aponte para a linha de Dívida Líquida no slide projetado e contraste o caixa positivo de ABEV3 com a dívida de R$ 21,3 bi de ASAI3.")

# 10. Conceito de EVA
add_subsection_header(doc, "4. O Conceito de Criação de Valor Econômico (EVA)")
add_meta_info(doc, "Slide 10 (Economic Value Added — Conceito)", "Explicar por que o lucro contábil não é o bastante se não cobrir o custo médio de capital (WACC).")
add_speech(doc,
    "Para compreendermos a real saúde financeira dessas empresas, aplicamos a metodologia clássica do Economic Value Added (EVA). "
    "Por que fizemos isso? Porque o lucro contábil tradicional que vemos na DRE ignora o custo do capital dos sócios. Uma empresa pode dar "
    "lucro contábil, mas se esse lucro for menor do que o retorno que o acionista teria em uma aplicação segura de renda fixa, na prática, "
    "a empresa está destruindo riqueza econômica. O EVA soluciona isso: ele deduz o custo médio ponderado de capital (WACC) do retorno real "
    "sobre o capital investido (ROIC pós-tax). O resultado dessa conta é o spread econômico. Se o spread for positivo (ROIC maior que WACC), "
    "a empresa gera valor líquido real. Se o spread for negativo (ROIC menor que WACC), a empresa destrói riqueza para o acionista."
)

# 11. Painel de Performance de Mercado
add_subsection_header(doc, "5. Painel de Performance e Risco Acionário Comparado (2021-2025)")
add_meta_info(doc, "Slide 11 (Painel de Performance de Mercado)", "Comparar as estatísticas consolidadas de mercado com o índice Ibovespa.")
add_speech(doc,
    "A tabela de performance e risco acionário sintetiza o comportamento das cotações diárias ajustadas ao longo de cinco anos, comparando "
    "diretamente os tickers ABEV3, ASAI3 e AMER3 com o benchmark do índice Ibovespa. O retorno acumulado do período mostra que o Ibovespa subiu "
    "46,03% no acumulado de 5 anos. Quando olhamos para a volatilidade anualizada, vemos a métrica clássica de desvio-padrão de retornos. "
    "O Ibovespa operou com volatilidade controlada de 17,46% ao ano. A Ambev se posicionou como um ativo muito próximo ao índice, com 23,34% "
    "de volatilidade. O Assaí, devido às pressões financeiras, viu sua volatilidade subir para 40,45% ao ano. E a Americanas oscilou de forma "
    "extrema, com desvio padrão anualizado de 129,45%. Ao analisarmos o Índice Sharpe — que mede o retorno excedente à taxa livre de risco "
    "(10,75%) por unidade de volatilidade —, observamos que todas as três ações apresentaram Sharpe negativo no período. Isso significa que, "
    "individualmente, nenhuma delas conseguiu bater a renda fixa sem que o investidor corresse riscos significativos no longo prazo."
)

# 12. Gráfico Comparativo
add_subsection_header(doc, "6. Análise da Curva Comparativa de Retornos")
add_meta_info(doc, "Slide 12 (Comparativo de Retornos Acumulados)", "Interpretar o gráfico comparativo de retornos acumulados históricos.")
add_speech(doc,
    "Este slide traz a representação visual dos retornos acumulados normalizados a partir da base zero de janeiro de 2021. Notem que a linha "
    "preta sólida do Ibovespa atua como a âncora do mercado financeiro, registrando uma alta de 46%. A linha azul da Ambev (ABEV3) demonstra "
    "um comportamento tipicamente defensivo. Ela não acompanha os ciclos de altas exponenciais da Bolsa, mas protege o capital nas quedas "
    "sistêmicas, fechando o período com retorno positivo de 29,28%. Em forte contraste, a linha amarela do Assaí (ASAI3) e a linha vermelha "
    "da Americanas (AMER3) se encontram abaixo da linha d'água de retorno zero. O Assaí acumula perda de 47,19%, evidenciando a correção de "
    "preços sofrida pela empresa pelo mercado diante do endividamento. E a linha de Americanas exibe o colapso patrimonial absoluto, com "
    "perda acumulada de 99,94%, despencando de forma abrupta logo após a revelação das inconsistências de lançamentos em janeiro de 2023."
)

# 13. Beta e Correlação
add_subsection_header(doc, "7. Beta Sistemático, Correlação e Diversificação de Carteira")
add_meta_info(doc, "Slide 13 (Beta, Correlação e Diversificação)", "Explicar a matriz de correlação cruzada e o risco sistemático (Beta).")
add_speech(doc,
    "Nesta seção do estudo, analisamos o risco sob a ótica da teoria moderna de carteiras de Harry Markowitz e do modelo CAPM de William Sharpe. "
    "Para isso, geramos a matriz de correlação de retornos diários e calculamos o coeficiente Beta sistemático contra o Ibovespa. O Beta mede "
    "a sensibilidade de um ativo às variações de mercado. O Ibovespa possui Beta padrão de 1,0. A matriz de correlação revela informações ricas "
    "para a montagem de portfólios. Reparem que a correlação da Ambev com o Ibovespa é de 0,48, e com a Americanas é de apenas 0,08. O Assaí tem "
    "correlação de 0,47 com o Ibovespa. Essas correlações moderadas a baixas provam que, em períodos normais, esses ativos reagem a fatores "
    "fundamentalistas específicos de microeconomia, e não apenas ao macro. Para um gestor de recursos, a Ambev surge como o principal "
    "diversificador strategic do grupo devido ao seu Beta de 0,64 e baixa correlação sistemática com os demais ativos."
)

# 14. Monte Carlo Teoria
add_subsection_header(doc, "8. A Simulação de Monte Carlo: Teoria e Parâmetros")
add_meta_info(doc, "Slide 14 (Simulação de Monte Carlo — Resultados)", "Apresentar a modelagem estocástica e a tabela de projeção de preços para 2026.")
add_speech(doc,
    "Para complementar a análise histórica do passado, desenvolvemos uma ferramenta de projeção para o futuro. Rodamos uma Simulação de "
    "Monte Carlo utilizando o modelo de Movimento Browniano Geométrico (MBG). Como essa simulação funciona? Ela calcula o retorno diário "
    "esperado das ações (a média ou drift histórico, representado pela letra grega mu) e o desvio padrão diário dos preços (a volatilidade, "
    "representada pela letra grega sigma) ao longo dos últimos 5 anos. A partir desses parâmetros, a fórmula gera caminhos estocásticos onde "
    "o preço de amanhã depende do preço de hoje mais uma componente de tendência histórica e um choque aleatório simulado baseado em uma "
    "distribuição normal padrão. Nós projetamos 10.000 caminhos possíveis para cada uma das três ações ao longo dos 252 dias úteis de 2026. "
    "Com esses 10.000 preços finais em dezembro de 2026, calculamos os percentis estatísticos e extraímos o VaR (Value at Risk) a 95% de "
    "confiança. O VaR nos diz qual a perda máxima esperada para o ativo no pior cenário de 5% de probabilidade."
)

# 15. Monte Carlo Funis
add_subsection_header(doc, "9. Diagnóstico Estocástico Comparado (Os Funis de Projeção)")
add_meta_info(doc, "Slide 15 (Simulação de Monte Carlo — Funis)", "Análise conceitual e visual dos três funis projetados em relação à volatilidade.")
add_speech(doc,
    "Os gráficos exibidos no slide representam os 'funis de projeção estocástica' resultantes da simulação de Monte Carlo. Cada gráfico mostra "
    "a dispersão das cotações em torno da mediana estatística. Visualmente, a largura de cada funil é o reflexo direto da volatilidade histórica "
    "de cada empresa. O funil da Ambev é estreito, simétrico e ascendente. Isso se deve à baixa volatilidade anualizada de 23,69% calibrada no script. "
    "A projeção de dispersão é contida, resultando em um VaR controlado de -30,27%. O funil da Americanas é bizarro, extremamente disperso e com "
    "forte viés negativo. A volatilidade absurda de 127,89% a.a. expande as caudas da distribuição de tal forma que a grande maioria dos 10.000 "
    "caminhos colapsa para próximo de zero, resultando em probabilidade de perda de 88,6%. O funil do Assaí exibe um comportamento intermediário. "
    "Ele se alarga significativamente, refletindo sua volatilidade de 40,45% a.a., com clara inclinação descendente da linha mediana que aponta "
    "para um ano de 2026 de correção, acumulando 63% de probabilidade de perda."
)

# 16. Conclusões e Recomendações
add_subsection_header(doc, "10. Conclusões Gerais e Recomendações de Investimento")
add_meta_info(doc, "Slide 16 (Conclusões e Recomendações)", "Encerrar a apresentação com as recomendações de alocação de portfólio.")
add_speech(doc,
    "Para finalizar o nosso estudo, sintetizamos quatro grandes lições práticas deste projeto multicritério. Primeira: A saúde contábil interna "
    "é o verdadeiro escudo do acionista. A Ambev prova que ter caixa líquido abundante e spread ROIC-WACC positivo protege o investidor contra "
    "drawdowns brutais na Bolsa. Segunda: A alavancagem financeira é uma faca de dois gumes. O Assaí tem um modelo de atacarejo excelente "
    "operacionalmente, mas a decisão de expandir agressivamente via dívidas bancárias sob juros de dois dígitos destruiu seu valor econômico "
    "residual (EVA negativo) e custou metade de seu valor de mercado. Terceira: Risco administrativo e governança são binários. A fraude de "
    "Americanas eliminou o patrimônio de forma permanente e irreversível, o que o Monte Carlo aponta como uma chance remota de recuperação acionária. "
    "Quarta: O Monte Carlo e o VaR estatístico são ferramentas indispensáveis de precificação de risco que devem caminhar de mãos dadas com a análise "
    "fundamentalista tradicional. Com isso, nossa recomendação é clara: priorizar alocação em ativos com spreads econômicos positivos e baixo "
    "endividamento em ciclos de juros altos, utilizando a volatilidade histórica para calcular a margem de segurança necessária no portfólio."
)

add_empty_line(doc)

# ────────────────────────────────────────────────────────────────────
# SEÇÃO 2: BLOCO TEMÁTICO — AMBEV S.A. (ABEV3)
# ────────────────────────────────────────────────────────────────────
add_section_header(doc, "Seção 2: Bloco Temático — Ambev S.A. (ABEV3)")

add_subsection_header(doc, "1. Perfil Corporativo e Modelo de Negócios")
add_meta_info(doc, "Slide 3 (Perfil Corporativo)", "Destacar a liderança de mercado e a capilaridade da distribuição da Ambev.")
add_speech(doc,
    "Vamos focar agora na análise detalhada da Ambev, sob o ticker ABEV3. A Ambev é a maior fabricante de cervejas e bebidas "
    "da América Latina, com uma participação de mercado dominante no Brasil. O grande trunfo da companhia reside no seu modelo "
    "de negócios baseado em escala e capilaridade. A empresa possui uma rede logística própria e exclusiva de distribuição "
    "física que atinge quase 100% dos bares, restaurantes e pontos de venda de bebidas do território nacional. Isso cria uma "
    "barreira de entrada intransponível para concorrentes de menor porte. Trata-se de um negócio clássico de consumo defensivo "
    "(ou não cíclico): mesmo em cenários de recessão econômica, a demanda por bebidas é altamente estável, o que confere à empresa "
    "uma fantástica previsibilidade de faturamento operacional de curto e longo prazo."
)

add_subsection_header(doc, "2. Estratégia de Mercado e Desafios (Premiumização e Digital)")
add_meta_info(doc, "Slide 4 (Atualidades e Estratégia + Gráfico)", "Analisar as novas frentes da Ambev e os gráficos de preços históricos.")
add_speech(doc,
    "Atualmente, a estratégia da Ambev está apoiada em dois grandes motores de crescimento. O primeiro é a Premiumização "
    "do portfólio. A empresa tem focado recursos de marketing e produção nas marcas de cervejas premium e super-premium, como "
    "Spaten, Corona e Stella Artois. Essa estratégia visa elevar o preço médio por hectolitro vendido e expandir a margem bruta de lucros. "
    "O segundo pilar é a digitalização. A Ambev criou o ecossistema digital 'Bees', uma plataforma B2B que revolucionou a forma como "
    "pequenos comerciantes fazem pedidos de bebidas (automatizando a cadeia e reduzindo custos), e o 'Zé Delivery', canal de delivery rápido "
    "que aproxima a marca diretamente do consumidor final. Como principais fatores de risco, destacam-se: as discussões sobre a Reforma "
    "Tributária, que pode criar impostos seletivos ('impostos do pecado') sobre bebidas alcoólicas, e potenciais mudanças regulatórias "
    "na dedutibilidade do Juros sobre Capital Próprio (JCP), que poderiam impactar a taxa de impostos efetiva da companhia no futuro."
)

add_subsection_header(doc, "3. Raio-X Contábil e de Balanço da Ambev (2025)")
add_meta_info(doc, "Slide 9 (Painel Comparativo Contábil)", "Demonstrar a solidez dos balanços e das margens de ABEV3 em 2025.")
add_speech(doc,
    "Ao abrirmos as demonstrações contábeis de 2025 da Ambev, os números comprovam sua posição de dominância financeira. A Receita Líquida "
    "atingiu 82,5 bilhões de reais, convertendo-se em um Lucro Bruto expressivo de 42,4 bilhões de reais, o que representa uma margem bruta "
    "elevada de 51,42%. Isso prova que mais de metade do faturamento líquido é retido após a dedução do custo de produção direta. "
    "O Lucro Operacional (EBIT) totalizou 21,9 bilhões de reais, gerando uma margem EBIT excelente de 26,54% — um patamar de eficiência operacional "
    "muito superior à média do varejo em geral. O Lucro Líquido consolidado foi de 14,9 bilhões de reais (margem líquida final de 18,12%). "
    "No Balanço Patrimonial, o Patrimônio Líquido de 82,5 bilhões de reais é suportado por uma estrutura de capital limpíssima: a empresa "
    "encerrou 2025 com R$ 8.440 milhões de Caixa Líquido positivo. Em termos práticos, a Ambev tem recursos em caixa suficientes para pagar "
    "todas as suas obrigações financeiras imediatas e ainda restam mais de 8,4 bilhões livres. Isso a torna imune a flutuações e juros altos."
)

add_subsection_header(doc, "4. Criação de Valor Econômico (EVA) da Ambev")
add_meta_info(doc, "Slide 10 (Economic Value Added — Ambev)", "Demonstrar a geração de riqueza e o spread ROIC-WACC positivo da Ambev.")
add_speech(doc,
    "Sob a ótica de geração de riqueza econômica, a Ambev é a única das três empresas analisadas que gerou riqueza líquida real em 2025. "
    "A companhia apresentou um Retorno sobre o Capital Investido (ROIC pós-tax) de 19,24%. O Custo Médio Ponderado de Capital (WACC) calculado "
    "para a empresa foi de 14,08%. O spread contábil pós-tax foi, portanto, de +5,16% positivo. Ao multiplicarmos esse spread positivo pela "
    "base de Capital Investido da empresa, chegamos a um EVA (Economic Value Added) positivo de mais de 4,1 bilhões de reais (precisamente "
    "R$ 4.148 milhões). Esse número expressa a riqueza excedente que a diretoria da Ambev gerou para o bolso de seus acionistas em 2025 após "
    "remunerar tanto o capital de terceiros quanto o capital próprio à taxa de retorno exigida pelo mercado de capitais."
)

add_subsection_header(doc, "5. Performance de Mercado e Risco das Ações da Ambev")
add_meta_info(doc, "Slide 11 (Painel de Performance) / Slide 12 (Gráfico)", "Explicar o retorno acumulado e a baixa volatilidade histórica da Ambev.")
add_speech(doc,
    "Ao avaliarmos as cotações diárias de ABEV3 ao longo dos 5 anos, percebemos que o comportamento na Bolsa reflete com precisão seus fundamentos "
    "contábeis. A ação acumulou retorno de 29,28% no período de 5 anos (CAGR anualizado de 5,27%). Embora tenha ficado abaixo dos 46% do Ibovespa, "
    "ela cumpriu perfeitamente o papel de ativo de proteção. A volatilidade anualizada de ABEV3 foi de apenas 23,34% — um patamar bastante controlado "
    "e muito inferior à dos pares do varejo. Seu Beta sistemático de 0,6430 comprova matematicamente o caráter defensivo do ativo: para cada 1% "
    "de variação no Ibovespa, a ação da Ambev oscila em média apenas 0,64%. Além disso, o Máximo Drawdown (pior queda histórica de pico a vale) "
    "foi de -36,18%, que ocorreu de forma suave ao longo de trimestres, mostrando que o ativo não sofre com pânicos repentinos de liquidez."
)

add_subsection_header(doc, "6. Projeções de Monte Carlo (2026) para ABEV3")
add_meta_info(doc, "Slide 14 (Tabela MC) / Slide 15 (Funil MC)", "Analisar as projeções estocásticas e o VaR da Ambev.")
add_speech(doc,
    "Quando projetamos o ano de 2026 para a Ambev via Simulação de Monte Carlo, os parâmetros estatísticos diários (mu de 0,000227 e volatilidade "
    "diária de 0,014925) resultaram no funil de preços mais estreito e controlado do grupo de estudo. Partindo do último preço de fechamento "
    "de R$ 13,86 em dezembro de 2025, a mediana projetada de 10.000 caminhos para dezembro de 2026 é de R$ 14,26, o que aponta para um crescimento "
    "estável esperado de +2,88%. A dispersão de risco é bastante contida: no pior cenário simulado (percentil de 5%), a ação cairia para R$ 9,67. "
    "Já no melhor cenário simulado (percentil de 95%), ela subiria para R$ 20,99. Isso resulta em um Value at Risk (VaR) de -30,27% de perda "
    "máxima projetada com 95% de confiança estatística. Além disso, a probabilidade da ação registrar qualquer perda nominal ao fim de 2026 "
    "é a menor do grupo: apenas 45,1%. O modelo estocástico reconhece a robustez de ABEV3 e aponta para um perfil de risco-retorno equilibrado."
)

add_empty_line(doc)

# ────────────────────────────────────────────────────────────────────
# SEÇÃO 3: BLOCO TEMÁTICO — ASSAÍ ATACADISTA (ASAI3)
# ────────────────────────────────────────────────────────────────────
add_section_header(doc, "Seção 3: Bloco Temático — Assaí Atacadista (ASAI3)")

add_subsection_header(doc, "1. Perfil Corporativo e Modelo de Negócios")
add_meta_info(doc, "Slide 5 (Perfil Corporativo)", "Destacar a dinâmica de alto giro e baixas margens do modelo de atacarejo do Assaí.")
add_speech(doc,
    "Iniciamos agora a análise específica do Assaí Atacadista, sob o ticker ASAI3. O Assaí atua no segmento de atacarejo, conhecido "
    "como Cash & Carry, que atende tanto clientes institucionais — como pequenos comércios, transformadores de alimentos e donos de "
    "restaurantes — quanto o consumidor final que busca economia em compras de grandes volumes. O modelo de negócios do atacarejo é baseado "
    "em margens operacionais muito estreitas e altíssimo giro de estoque. A vantagem competitiva do Assaí reside na sua capacidade de negociar "
    "volumes imensos com as indústrias fornecedoras e repassar essa economia em preços baixos na gôndola. Em termos de histórico societário, "
    "o Assaí pertencia ao Grupo Pão de Açúcar (GPA) e realizou um bem-sucedido processo de spin-off (cisão corporativa pura) concluído em 2021, "
    "tornando-se uma corporação pulverizada e focada exclusivamente em atacarejo alimentar."
)

add_subsection_header(doc, "2. Expansão, Alavancagem e Gráfico de Preços do Assaí")
add_meta_info(doc, "Slide 6 (Atualidades e Estratégia + Gráfico)", "Analisar as conversões das lojas Extra Hiper e a pressão das dívidas.")
add_speech(doc,
    "Atualmente, a tese do Assaí tem sido marcada por uma das maiores e mais agressivas expansões geográficas do varejo nacional. "
    "O marco dessa estratégia foi a aquisição e conversão de 70 pontos de venda do antigo Extra Hipermercado. Essa transação permitiu ao "
    "Assaí ocupar espaços comerciais nobres dentro de grandes centros urbanos de forma rápida. Porém, para financiar essa compra de pontos "
    "comerciais e arcar com o Capex bilionário de conversão das lojas em atacarejo, o Assaí captou um volume massivo de empréstimos bancários. "
    "A dívida bruta saltou para patamares elevados. A grande tempestade estratégica ocorreu quando a taxa de juros Selic subiu e permaneceu "
    "alta a dois dígitos por um período prolongado. A empresa viu as despesas de juros devorarem a sua receita operacional. Embora a receita "
    "bruta tenha crescido com a abertura das lojas Extra convertidas, o custo de carregar essa dívida líquida bilionária pressionou fortemente "
    "a última linha da DRE, deprimindo o lucro líquido acionário."
)

add_subsection_header(doc, "3. Raio-X Contábil e de Balanço do Assaí (2025)")
add_meta_info(doc, "Slide 9 (Painel Comparativo Contábil)", "Demonstrar a receita volumosa contraída com margem operacional enxuta de ASAI3.")
add_speech(doc,
    "Ao detalharmos os demonstrativos contábeis de 2025 do Assaí, o volume de vendas impressiona: a Receita Líquida foi de 71,5 bilhões "
    "de reais (muito próxima do faturamento da Ambev). O Lucro Bruto atingiu 12 bilhões de reais. Devido ao modelo de baixo markup do atacarejo, "
    "a margem bruta de lucros do Assaí é de apenas 16,87% (menos de um terço da margem de Ambev). O EBITDA fechou em 3,6 bilhões de reais "
    "(margem EBITDA de 5,07%). O Lucro Operacional (EBIT) encerrou 2025 também em 3,6 bilhões de reais, indicando uma margem operacional "
    "consolidada de 5,07%. E o Lucro Líquido final foi de 1 bilhão de reais, resultando em uma margem líquida enxuta de 1,43%. A grande "
    "vulnerabilidade estrutural reside no Balanço Patrimonial. O Patrimônio Líquido do Assaí de 4,7 bilhões de reais suporta uma Dívida Líquida "
    "volumosa de 21,3 bilhões de reais. Toda essa dívida consome o caixa gerado no fluxo operacional com despesas financeiras pesadas."
)

add_subsection_header(doc, "4. Destruição de Valor Econômico (EVA) no Assaí")
add_meta_info(doc, "Slide 10 (Economic Value Added — Assaí)", "Explicar a destruição de valor residual sob o custo do capital investido do Assaí.")
add_speech(doc,
    "O caso do Assaí é um exemplo clássico da diferença entre o lucro contábil e o lucro econômico medido pelo EVA. Contabilmente, em 2025, "
    "o Assaí registrou um lucro líquido de 1 bilhão de reais na DRE. No entanto, sob a ótica econômica, a empresa destruiu riqueza do acionista. "
    "O Retorno sobre o Capital Investido (ROIC pós-tax) do Assaí foi de 8,89%. O Custo Médio Ponderado de Capital (WACC) calculado para carregar "
    "a empresa sob juros altos foi de 10,14%. O spread contábil pós-tax foi negativo: de -1,25%. Isso significa que o retorno gerado pelas "
    "operações do Assaí foi menor do que o custo de carregar a estrutura de capital (dívida + capital próprio). Multiplicando esse spread "
    "negativo pela base de capital investido de mais de 26 bilhões de reais, o Assaí registrou um EVA negativo de -333 milhões de reais em 2025."
)

add_subsection_header(doc, "5. Performance de Mercado e Risco das Ações do Assaí")
add_meta_info(doc, "Slide 11 (Painel de Performance) / Slide 12 (Gráfico)", "Explicar a desvalorização na Bolsa e a sensibilidade do papel do Assaí aos juros.")
add_speech(doc,
    "A performance da ação ASAI3 na Bolsa mostra que o mercado de capitais antecipou e precificou a destruição de valor econômico da empresa. "
    "A ação acumulou uma desvalorização de -47,19% no período de 5 anos (uma queda média anualizada CAGR de -11,99%). A volatilidade anualizada "
    "de ASAI3 foi de 40,45% — mais que o dobro da volatilidade do Ibovespa. Seu Beta de 1,0981 indica que a ação oscila alinhada ao mercado de ações "
    "em geral, mas amplifica os movimentos de queda quando o mercado sofre. O Máximo Drawdown (queda máxima sofrida de pico a vale na Bolsa) foi "
    "devastador: -75,55% de desvalorização em relação ao seu topo histórico acionário. A forte correlação do papel com a taxa de juros futura (DI) "
    "torna a cotação do Assaí extremamente sensível aos comunicados do Copom sobre a manutenção da taxa Selic elevada."
)

add_subsection_header(doc, "6. Projeções de Monte Carlo (2026) para ASAI3")
add_meta_info(doc, "Slide 14 (Tabela MC) / Slide 15 (Funil MC)", "Analisar a distribuição estocástica de preços e o VaR do Assaí.")
add_speech(doc,
    "Ao rodarmos a Simulação de Monte Carlo para o Assaí em 2026, os parâmetros de tendência e volatilidade (mu de -0,000205 e volatilidade "
    "diária de 0,025483) projetam um funil de preços intermediário, mas com nítido viés de queda. Partindo do preço real de R$ 7,20 em dezembro "
    "de 2025, o modelo projeta uma cotação mediana de R$ 6,30 para dezembro de 2026, o que representaria uma desvalorização estimada de -12,54% "
    "ao longo do ano. A dispersão de risco é relevante: no pior cenário simulado (percentil de 5%), o preço desaba para R$ 3,23. No melhor cenário "
    "(percentil de 95%), ele sobe para R$ 12,26. Esse comportamento projeta um Value at Risk (VaR) de -55,09% com 95% de confiança. Além disso, "
    "a probabilidade do Assaí fechar o ano de 2026 acumulando perda nominal é de expressivos 63,0%. Isso alerta os investidores sobre o risco do papel."
)

add_empty_line(doc)

# ────────────────────────────────────────────────────────────────────
# SEÇÃO 4: BLOCO TEMÁTICO — LOJAS AMERICANAS (AMER3)
# ────────────────────────────────────────────────────────────────────
add_section_header(doc, "Seção 4: Bloco Temático — Lojas Americanas (AMER3)")

add_subsection_header(doc, "1. Perfil Corporativo e Modelo de Negócios")
add_meta_info(doc, "Slide 7 (Perfil Corporativo)", "Analisar o histórico da rede física e a operação multicanal física-digital da Americanas.")
add_speech(doc,
    "Passamos agora para a análise detalhada das Lojas Americanas, sob o ticker AMER3. Fundada em 1929, as Americanas constituem uma das marcas "
    "mais tradicionais do varejo brasileiro, com forte presença física em shoppings e centros urbanos. O modelo de negócios da empresa é baseado "
    "em conveniência e multicanalidade. Ela integrava uma capilarizada rede de lojas físicas de varejo de conveniência de baixo ticket com uma "
    "robusta plataforma digital de e-commerce e marketplace de grande porte. O foco comercial histórico da empresa sempre foi a venda de bens "
    "de consumo de giro rápido, chocolates, brinquedos e eletrodomésticos. Contudo, a tese de crescimento do e-commerce digital no Brasil exigia "
    "forte subsídio de preços e capital intensivo, expondo as fraquezas estruturais de concorrência com players globais e dependência de crédito."
)

add_subsection_header(doc, "2. A Fraude Contábil e o Processo de Recuperação Judicial")
add_meta_info(doc, "Slide 8 (Atualidades e Estratégia + Gráfico)", "Analisar o rombo contábil de R$ 20 bilhões, a RJ e o grupamento de ações.")
add_speech(doc,
    "A história das Americanas mudou radicalmente em 11 de janeiro de 2023 com a divulgação do histórico fato relevante informando a descoberta "
    "de 'inconsistências contábeis' estimadas preliminarmente em R$ 20 bilhões. Investigações posteriores revelaram tratar-se de fraude contábil "
    "intencional que omitia do passivo as operações de risco sacado (forfaiting) para inflar artificialmente o caixa e os lucros. O colapso foi imediato. "
    "A empresa perdeu acesso a crédito, as linhas bancárias foram congeladas, e fornecedores cortaram o abastecimento. Em poucos dias, a Americanas "
    "entrou com pedido de Recuperação Judicial urgente. Para tentar salvar a operação, um plano de reestruturação foi aprovado prevendo um aumento "
    "de capital de R$ 12 bilhões com recursos dos acionistas de referência e a conversão de parte significativa das dívidas dos bancos credores em ações. "
    "Para completar, em 26 de agosto de 2024, a B3 executou um grupamento de ações na proporção de 100 para 1, para retirar a ação da zona de penny stock. "
    "Esse grupamento explica as cotações passadas de milhares de reais que aparecem no gráfico histórico ajustado."
)

add_subsection_header(doc, "3. Raio-X Contábil e de Balanço de Americanas (2025)")
add_meta_info(doc, "Slide 9 (Painel Comparativo Contábil)", "Demonstrar a contração severa da receita e o prejuízo reportado de AMER3 em 2025.")
add_speech(doc,
    "O balanço contábil de 2025 da Americanas é o retrato físico de uma corporação em ruínas operacionais. A Receita Líquida encolheu para "
    "apenas 12,3 bilhões de reais (uma fração do que já faturou no passado). O Lucro Bruto foi de 3,1 bilhões de reais, indicando margem bruta "
    "de 25,34%. O EBITDA foi de modestos 684 milhões de reais (margem EBITDA de 5,55%), que mal cobre as despesas básicas da estrutura física. "
    "O Lucro Operacional (EBIT) foi de apenas 331 milhões de reais, com margem operacional deprimida de 2,69%. A empresa registrou prejuízo "
    "líquido contábil consolidado de 271 milhões de reais na última linha. No Balanço Patrimonial, a Americanas reporta um Patrimônio Líquido "
    "de 5,3 bilhões de reais, contudo, esse PL só existe no papel após as injeções recentes de capital dos sócios. A Dívida Líquida reportada "
    "é de 5,7 bilhões de reais (uma dívida bruta de 9,3 bilhões contra 3,5 bilhões de caixa). Os indicadores apontam para uma operação contraída."
)

add_subsection_header(doc, "4. Destruição de Valor Econômico (EVA) em Americanas")
add_meta_info(doc, "Slide 10 (Economic Value Added — Americanas)", "Analisar o spread ROIC-WACC altamente negativo e o EVA destruído.")
add_speech(doc,
    "Ao avaliarmos o Economic Value Added (EVA) da Americanas, os números revelam uma destruição de riqueza sem precedentes. O Retorno sobre "
    "o Capital Investido (ROIC pós-tax) da Americanas foi de apenas 2,07%. Enquanto isso, o WACC (Custo Médio Ponderado de Capital) exigido para "
    "carregar a empresa devido ao risco de insolvência saltou para 12,28%. O spread contábil pós-tax foi negativo de catastróficos -10,21%. "
    "Isso significa que para cada real investido na operação das Americanas, o investidor perdeu mais de 10 centavos em relação à taxa de remuneração "
    "mínima aceitável. O EVA final foi negativo de -1.074 milhões de reais em 2025. Esse dado contesta a viabilidade econômica do modelo atual "
    "e evidencia a destruição severa de patrimônio líquido provocada pela fraude e pelo endividamento estrutural da Recuperação Judicial."
)

add_subsection_header(doc, "5. Performance de Mercado e Risco das Ações de Americanas")
add_meta_info(doc, "Slide 11 (Painel de Performance) / Slide 12 (Gráfico)", "Explicar o colapso patrimonial de -99,94% nas cotações e a volatilidade extrema.")
add_speech(doc,
    "O comportamento de AMER3 na Bolsa de Valores representa o colapso patrimonial mais rápido da história do mercado de capitais brasileiro. "
    "O retorno acumulado do período de 5 anos foi negativo em absurdos -99,94% (uma taxa geométrica anualizada de queda de -76,96%). Em termos "
    "práticos, quem investiu R$ 10.000 em AMER3 em janeiro de 2021 encerrou 2025 com apenas R$ 6,00. A volatilidade anualizada foi a maior registrada "
    "no mercado, fechando em 129,45%, refletindo movimentos diários de dezenas de por cento de alta ou queda sem sustentação de fundamentos. "
    "O Beta de 1,5005 indica uma volatilidade sistemática exagerada. O Máximo Drawdown alcançou -99,96%. O gráfico de preços ajustado por grupamento "
    "de cotações registra a descida ao inferno financeiro: saindo de um patamar equivalente a R$ 7.354,00 por ação em 2021 para fechar 2025 a R$ 5,13."
)

add_subsection_header(doc, "6. Projeções de Monte Carlo (2026) para AMER3")
add_meta_info(doc, "Slide 14 (Tabela MC) / Slide 15 (Funil MC)", "Demonstrar a distribuição estocástica de cotações para 2026, com 88,6% de probabilidade de perda.")
add_speech(doc,
    "Ao rodarmos a Simulação de Monte Carlo para a Americanas para o ano de 2026, os parâmetros históricos catastróficos (mu diário negativo de "
    "-0,002808 e volatilidade diária de 0,080562) produzem um funil estocástico dramático e de altíssima dispersão de caudas. Partindo do último "
    "preço de fechamento de R$ 5,13 em dezembro de 2025, o preço mediano projetado para dezembro de 2026 é de R$ 1,11, apontando para uma desvalorização "
    "esperada de -78,38% adicionais a partir do patamar de preço atual. O risco estatístico é máximo: o Value at Risk (VaR) a 95% de confiança "
    "aponta para uma perda de -97,26% (onde a ação cairia para míseros R$ 0,14 no pior cenário). O melhor cenário projeta R$ 9,05, mostrando que a "
    "volatilidade cria uma falsa impressão de ganhos elevados em cenários raros. A probabilidade da Americanas fechar 2026 gerando perda real "
    "é de 88,6%. O modelo estocástico confirma que o papel opera em termos puramente especulativos, com alta probabilidade de destruição acionária."
)

# Rodapé final centralizado
add_empty_line(doc)
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run("Fim do Roteiro de Falas Oficial — Análise de Desempenho e Risco (2021-2025)")
run_footer.italic = True
run_footer.font.name = FONT_NAME
run_footer.font.size = Pt(9.5)
run_footer.font.color.rgb = GRAY_TEXT

# Salva o arquivo de saída
output_path = r"c:\andd\ap2_analise_de_dados\3_Roteiro_de_Falas_Speech.docx"
doc.save(output_path)
print(f"Roteiro DOCX salvo com sucesso em: {output_path}")
